"""审计报告导出:TXT / Word(带批注+高亮) / PDF(中文字体)。

对应论文 4.3.4/5.4。
- Word:python-docx 1.2+ 的 add_comment 添加真批注;高/中/低风险条款用颜色+背景标注
- PDF:reportlab 注册系统中文字体(SimHei),支持中文输出
"""
from __future__ import annotations

import os
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.lib import colors


# ===== 字体注册(系统中文字体,Windows) =====
_FONT_REGISTERED = False
_FONT_NAME = "ChineseFont"


def _ensure_font() -> str:
    global _FONT_REGISTERED
    if _FONT_REGISTERED:
        return _FONT_NAME
    candidates = [
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for path in candidates:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont(_FONT_NAME, path))
                _FONT_REGISTERED = True
                return _FONT_NAME
            except Exception:
                continue
    # 退路:使用内建字体
    _FONT_REGISTERED = True
    return "Helvetica"


# ===== 数据结构:一次审计的整理结果 =====
def _level_to_color(level: str) -> RGBColor:
    """风险等级 → Word 字体颜色。"""
    if level == "高风险" or level == "高":
        return RGBColor(0xC0, 0x39, 0x2B)
    if level == "中风险" or level == "中":
        return RGBColor(0xCC, 0x99, 0x00)
    return RGBColor(0x27, 0xAE, 0x60)


def _level_to_highlight(level: str) -> Optional[int]:
    if level == "高风险" or level == "高":
        return WD_COLOR_INDEX.RED
    if level == "中风险" or level == "中":
        return WD_COLOR_INDEX.YELLOW
    return WD_COLOR_INDEX.BRIGHT_GREEN


# ============ TXT ============
def export_txt(audit: dict, out_path: str) -> str:
    lines = []
    lines.append("=" * 50)
    lines.append("合同条款智能审计报告")
    lines.append("=" * 50)
    lines.append("")
    lines.append(f"【待审条款】\n{audit.get('clause', '')}")
    lines.append("")
    lines.append(f"【相关法规】\n{audit.get('law_context', '无')}")
    lines.append("")
    if audit.get("historical_context"):
        lines.append(f"【历史采纳建议】\n{audit['historical_context']}")
        lines.append("")
    lines.append(f"【条款类型】{audit.get('risk_type', '其他')}")
    lines.append(f"【风险等级】{audit.get('level', '中风险')}")
    lines.append(f"【风险评分】{audit.get('formula', '')}")
    lines.append(f"【LLM 置信度】{audit.get('confidence', 0)}")
    lines.append(f"【需人工复核】{'是' if audit.get('need_review') else '否'}")
    lines.append("")
    lines.append(f"【风险分析】\n{audit.get('reason', '')}")
    lines.append("")
    lines.append(f"【修改建议】\n{audit.get('suggestion', '')}")
    lines.append("")
    if audit.get("rule_hits"):
        lines.append("【规则引擎命中】")
        for h in audit["rule_hits"]:
            lines.append(f"- {h.get('rule_id', '')} [{h.get('risk_level', '')}] {h.get('reason', '')}")
    lines.append("")
    lines.append(f"【计时】{audit.get('timing', {})}")
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return out_path


# ============ DOCX (with comments) ============
def export_docx(audit: dict, out_path: str) -> str:
    doc = Document()
    title = doc.add_heading("合同条款智能审计报告", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    # 摘要
    doc.add_heading("一、审计摘要", level=1)
    summary_tbl = doc.add_table(rows=4, cols=2)
    summary_tbl.style = "Light Grid Accent 1"
    rows = [
        ("条款类型", audit.get("risk_type", "其他")),
        ("风险等级", audit.get("level", "中风险")),
        ("风险评分", audit.get("formula", "")),
        ("是否需要复核", "是" if audit.get("need_review") else "否"),
    ]
    for i, (k, v) in enumerate(rows):
        summary_tbl.cell(i, 0).text = k
        summary_tbl.cell(i, 1).text = str(v)

    # 待审条款(带批注)
    doc.add_heading("二、待审条款", level=1)
    p = doc.add_paragraph()
    run = p.add_run(audit.get("clause", ""))
    run.font.size = Pt(10.5)
    hl = _level_to_highlight(audit.get("level", "中风险"))
    if hl is not None:
        run.font.highlight_color = hl

    # 加 Word 真批注
    comment_text = (
        f"【{audit.get('level', '中风险')}】{audit.get('reason', '')}\n"
        f"建议:{audit.get('suggestion', '')}"
    )
    try:
        doc.add_comment(runs=[run], text=comment_text, author="智能审计系统", initials="AI")
    except Exception:
        # 极端情况退化为段落
        warn = doc.add_paragraph()
        warn_run = warn.add_run(f"【批注】{comment_text}")
        warn_run.font.color.rgb = _level_to_color(audit.get("level", "中风险"))

    # 法规引用
    doc.add_heading("三、法规依据", level=1)
    doc.add_paragraph(audit.get("law_context", "无") or "无")

    # 历史参考
    if audit.get("historical_context"):
        doc.add_heading("四、历史采纳建议参考", level=1)
        doc.add_paragraph(audit["historical_context"])

    # 详细分析
    doc.add_heading("五、详细分析", level=1)
    p_reason = doc.add_paragraph()
    p_reason.add_run("分析:").bold = True
    p_reason.add_run(audit.get("reason", ""))

    p_sug = doc.add_paragraph()
    p_sug.add_run("修改建议:").bold = True
    sug_run = p_sug.add_run(audit.get("suggestion", ""))
    sug_run.font.color.rgb = _level_to_color(audit.get("level", "中风险"))

    # 规则命中
    if audit.get("rule_hits"):
        doc.add_heading("六、规则引擎命中", level=1)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Light Grid Accent 1"
        hdr = tbl.rows[0].cells
        hdr[0].text = "规则ID"
        hdr[1].text = "条款类型"
        hdr[2].text = "等级"
        hdr[3].text = "依据"
        for h in audit["rule_hits"]:
            row = tbl.add_row().cells
            row[0].text = h.get("rule_id", "")
            row[1].text = h.get("risk_type", "")
            row[2].text = h.get("risk_level", "")
            row[3].text = h.get("reason", "")

    doc.add_heading("七、性能耗时", level=1)
    timing = audit.get("timing", {}) or {}
    if timing:
        t = doc.add_table(rows=1, cols=2)
        t.style = "Light Grid Accent 1"
        t.rows[0].cells[0].text = "阶段"
        t.rows[0].cells[1].text = "耗时(秒)"
        for k, v in timing.items():
            row = t.add_row().cells
            row[0].text = str(k)
            row[1].text = f"{v:.3f}" if isinstance(v, (int, float)) else str(v)

    doc.save(out_path)
    return out_path


# ============ PDF (Chinese support) ============
def export_pdf(audit: dict, out_path: str) -> str:
    font = _ensure_font()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleCN", parent=styles["Title"], fontName=font, fontSize=18, leading=22, spaceAfter=12
    )
    h_style = ParagraphStyle(
        "HeadingCN", parent=styles["Heading2"], fontName=font, fontSize=13, leading=18, spaceAfter=6, textColor=colors.HexColor("#2C3E50")
    )
    body_style = ParagraphStyle(
        "BodyCN", parent=styles["BodyText"], fontName=font, fontSize=10.5, leading=16, spaceAfter=4
    )

    doc = SimpleDocTemplate(
        out_path, pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
    )
    flow = []
    flow.append(Paragraph("合同条款智能审计报告", title_style))
    flow.append(Spacer(1, 6))

    flow.append(Paragraph("一、审计摘要", h_style))
    summary = [
        ["条款类型", audit.get("risk_type", "其他")],
        ["风险等级", audit.get("level", "中风险")],
        ["风险评分", audit.get("formula", "")],
        ["是否复核", "是" if audit.get("need_review") else "否"],
        ["LLM置信度", str(audit.get("confidence", 0))],
    ]
    tbl = Table(summary, colWidths=[3.5 * cm, 13 * cm])
    tbl.setStyle(TableStyle([
        ("FONT", (0, 0), (-1, -1), font, 10),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CCCCCC")),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F0F4F8")),
    ]))
    flow.append(tbl)
    flow.append(Spacer(1, 10))

    flow.append(Paragraph("二、待审条款", h_style))
    flow.append(Paragraph(_escape(audit.get("clause", "")), body_style))

    flow.append(Paragraph("三、法规依据", h_style))
    flow.append(Paragraph(_escape(audit.get("law_context", "") or "无"), body_style))

    if audit.get("historical_context"):
        flow.append(Paragraph("四、历史采纳建议参考", h_style))
        flow.append(Paragraph(_escape(audit["historical_context"]), body_style))

    flow.append(Paragraph("五、详细分析", h_style))
    flow.append(Paragraph(f"<b>分析:</b>{_escape(audit.get('reason', ''))}", body_style))
    flow.append(Paragraph(f"<b>修改建议:</b>{_escape(audit.get('suggestion', ''))}", body_style))

    if audit.get("rule_hits"):
        flow.append(Paragraph("六、规则引擎命中", h_style))
        data = [["规则ID", "条款类型", "等级", "依据"]]
        for h in audit["rule_hits"]:
            data.append([
                h.get("rule_id", ""),
                h.get("risk_type", ""),
                h.get("risk_level", ""),
                _shorten(h.get("reason", ""), 50),
            ])
        rt = Table(data, colWidths=[2 * cm, 2.5 * cm, 1.5 * cm, 10.5 * cm])
        rt.setStyle(TableStyle([
            ("FONT", (0, 0), (-1, -1), font, 9.5),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CCCCCC")),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F0F4F8")),
        ]))
        flow.append(rt)

    timing = audit.get("timing", {}) or {}
    if timing:
        flow.append(Spacer(1, 8))
        flow.append(Paragraph("七、性能耗时", h_style))
        rows = [["阶段", "耗时(秒)"]]
        for k, v in timing.items():
            rows.append([str(k), f"{v:.3f}" if isinstance(v, (int, float)) else str(v)])
        tt = Table(rows, colWidths=[6 * cm, 4 * cm])
        tt.setStyle(TableStyle([
            ("FONT", (0, 0), (-1, -1), font, 10),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CCCCCC")),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F0F4F8")),
        ]))
        flow.append(tt)

    doc.build(flow)
    return out_path


def _escape(text: str) -> str:
    if not text:
        return ""
    return (
        text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")
    )


def _shorten(text: str, n: int) -> str:
    if len(text) <= n:
        return text
    return text[: n - 1] + "…"
