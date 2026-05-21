"""文件解析:PyMuPDF + python-docx + 纯文本。

对应论文 4.3.1 节,但按用户决定不接入 PaddleOCR(仅文字版PDF)。
扫描件检测:若 PyMuPDF 抽取后正文极短,提示用户'扫描件需OCR,当前未启用'。
"""
from __future__ import annotations

import io
from typing import Optional, Tuple

import fitz  # PyMuPDF
from docx import Document


SCAN_HINT = "[扫描件检测] 该 PDF 未抽到足够文本,可能为扫描件。当前未启用 PaddleOCR,请改用文本版 PDF/Word 或 TXT。"


def parse_txt(file_bytes: bytes) -> Tuple[Optional[str], Optional[str]]:
    try:
        for enc in ("utf-8", "gbk", "gb18030", "latin-1"):
            try:
                return file_bytes.decode(enc), None
            except UnicodeDecodeError:
                continue
        return None, "TXT 编码无法识别"
    except Exception as e:
        return None, f"TXT 解析失败:{e}"


def parse_docx(file_bytes: bytes) -> Tuple[Optional[str], Optional[str]]:
    try:
        doc = Document(io.BytesIO(file_bytes))
        seen = set()
        parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        # 表格中的文本也抽出来（去重，避免嵌套表格重复）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    txt = cell.text.strip() if cell.text else ""
                    if txt and txt not in seen:
                        seen.add(txt)
                        parts.append(txt)
        text = "\n".join(parts).strip()
        if not text:
            return None, "DOCX 内容为空"
        return text, None
    except Exception as e:
        return None, f"DOCX 解析失败:{e}"


def parse_pdf(file_bytes: bytes) -> Tuple[Optional[str], Optional[str]]:
    """用 PyMuPDF 抽取,文本极少时给出扫描件提示。"""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        parts = []
        for page in doc:
            txt = page.get_text("text")
            if txt:
                parts.append(txt)
        doc.close()
        text = "\n".join(parts).strip()
        if not text:
            return None, SCAN_HINT
        # 检测:平均每页 < 30 字符基本是图片版
        avg = len(text) / max(1, len(parts))
        if avg < 30 and len(text) < 200:
            return text, SCAN_HINT
        return text, None
    except Exception as e:
        return None, f"PDF 解析失败:{e}"


def parse_upload(filename: str, file_bytes: bytes) -> Tuple[Optional[str], Optional[str]]:
    name = (filename or "").lower()
    if name.endswith(".txt"):
        return parse_txt(file_bytes)
    if name.endswith(".docx"):
        return parse_docx(file_bytes)
    if name.endswith(".pdf"):
        return parse_pdf(file_bytes)
    return None, "不支持的文件格式,请上传 .txt / .docx / .pdf"
